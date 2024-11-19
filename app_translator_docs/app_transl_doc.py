import requests
import os
from docx import Document

subscription_key = os.getenv("AZURE_API_KEY")
endpoint = "https://api.cognitive.microsofttranslator.com/"
location = "eastus2"
language_destination = "pt-br"

def translator_text(text, target_language):
    path = '/translate'
    constructed_url = endpoint + path
    headers = {
        'Ocp-Apim-Subscription-Key': subscription_key,
        'Ocp-Apim-Subscription-Region': location,
        'Content-type': 'application/json',
        'X-ClientTraceId': str(os.urandom(16))
    }

    body = [{'text': text}]
    params = {'api-version': '3.0', 'from': 'en', 'to': target_language}

    request = requests.post(constructed_url, params=params, headers=headers, json=body)
    response = request.json()

    try:
        return response[0]["translations"][0]["text"]
    except (KeyError, IndexError):
        print(f"Error in response: {response}")
        return "[Translation Error]"

def translate_document(path):
    document = Document(path)
    full_text = []

    # Iterate through paragraphs and trnaslate 
    for paragraph in document.paragraphs:
        # Ignore empty paragraphs 
        if paragraph.text.strip():  
            translated_text = translator_text(paragraph.text, language_destination)
            full_text.append(translated_text)
        else:
            # Add a empty line if the original paragraph is empty
            full_text.append("")  

    # Create a new document and add translated paragraphs
    translated_doc = Document()
    for line in full_text:
        translated_doc.add_paragraph(line)

    # Save translated document
    path_translated = path.replace(".docx", f"_{language_destination}.docx")
    translated_doc.save(path_translated)
    return path_translated

# Translate a document
input_file = "taylor_swift_song.docx"
output_file = translate_document(input_file)
print(f"Translated document saved to: {output_file}")
